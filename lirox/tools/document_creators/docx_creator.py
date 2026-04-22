"""DOCX Creator — Rich Word Document Engine.

Produces a .docx file with:
  - Topic-aware primary color applied to headings
  - Cover page (title, author, date)
  - Section headings, body paragraphs, bullet lists
  - Optional inline tables (via ``sec["table"]``)
  - Page break between cover and content
"""
from __future__ import annotations

import logging
import os
from typing import Any, Dict, List

from lirox.verify import FileReceipt
from lirox.tools.document_creators.base import ensure_dep, pick_palette, PALETTES
from lirox.safety.audit import log_audit_event

_logger = logging.getLogger("lirox.document_creators.docx")


def create_docx(path: str, title: str, sections: List[Dict[str, Any]],
                query: str = "", user_name: str = "") -> FileReceipt:
    """Create a Word document with headings, paragraphs, bullets, and tables.

    Parameters
    ----------
    path:       Absolute output path.
    title:      Document title.
    sections:   List of dicts with keys ``heading``, ``body``, ``bullets``,
                and optionally ``table`` (list of rows, each a list of values).
    query:      Original user request — used for palette selection.
    user_name:  Shown as author on the cover page.

    Returns
    -------
    FileReceipt with ``ok=True`` and ``verified=True`` on success.
    """
    from pathlib import Path as _Path
    out_path = _Path(path).resolve()
    r = FileReceipt(tool="docx_creator", operation="create_docx", path=str(out_path))

    # ── PRE-FLIGHT VALIDATION (v1.1 fix) ──
    if not path:
        r.error = "Path cannot be empty"
        return r
    if not title:
        title = "Untitled Document"
    if user_name is None:
        user_name = ""
    if sections is None or not isinstance(sections, list):
        sections = [{"heading": title, "body": "", "bullets": []}]
    if not sections:
        sections = [{"heading": title, "body": "", "bullets": []}]

    try:
        ensure_dep("python-docx", "docx")
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        out_dir = out_path.parent
        out_dir.mkdir(parents=True, exist_ok=True)
        if not os.access(str(out_dir), os.W_OK):
            r.error = f"Output directory is not writable: {out_dir}"
            return r

        palette_name = pick_palette(query or title, title)
        pal = PALETTES[palette_name]
        primary_rgb = RGBColor(
            int(pal["primary"][:2], 16),
            int(pal["primary"][2:4], 16),
            int(pal["primary"][4:6], 16),
        )

        doc = Document()

        # Cover: title
        t = doc.add_heading(title, level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.color.rgb = primary_rgb

        # Cover: author
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_text = f"Prepared for: {user_name}" if user_name else "Generated Document"
        run = p.add_run(author_text)
        run.font.size = Pt(12)
        run.font.italic = True

        # Cover: date
        from datetime import datetime
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(10)

        doc.add_page_break()

        # Content sections
        for sec in sections:
            if sec.get("heading"):
                h = doc.add_heading(sec["heading"], level=1)
                for run in h.runs:
                    run.font.color.rgb = primary_rgb

            if sec.get("body"):
                for para in sec["body"].split("\n\n"):
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.paragraph_format.space_after = Pt(8)

            if sec.get("bullets"):
                for b in sec["bullets"]:
                    if b and b.strip():
                        doc.add_paragraph(b.strip(), style="List Bullet")

            if sec.get("table"):
                data = sec["table"]
                if data and len(data) > 0:
                    table = doc.add_table(rows=len(data), cols=len(data[0]))
                    table.style = "Light Grid Accent 1"
                    for i, row_data in enumerate(data):
                        for j, cell_val in enumerate(row_data):
                            table.cell(i, j).text = str(cell_val)

        doc.save(str(out_path))

        if out_path.exists():
            r.ok = True
            r.verified = True
            r.bytes_written = out_path.stat().st_size
            r.message = (
                f"Created Word doc: {out_path} "
                f"({r.bytes_written:,} bytes, {len(sections)} sections)"
            )
            r.details["section_count"] = len(sections)
            r.details["palette"] = palette_name
        else:
            r.error = "Word build completed but file not found on disk"
        log_audit_event(
            "document_create_docx",
            str(out_path),
            status="ok" if (r.ok and r.verified) else "error",
            detail=r.message or r.error,
        )
        return r
    except Exception as e:
        r.error = f"Word creation error: {e}"
        log_audit_event("document_create_docx", str(out_path), status="error", detail=r.error)
        return r
